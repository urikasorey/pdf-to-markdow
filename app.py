import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io
import os
from pathlib import Path
import base64
import re
import cv2
import numpy as np
import pandas as pd
import zipfile
from datetime import datetime
import time
try:
    import pytesseract
    # Streamlit Cloud sẽ cài Tesseract tự động
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

def optimize_image(image_path, max_width=1200, quality=85):
    """Tối ưu hóa kích thước và chất lượng hình ảnh"""
    try:
        img = Image.open(image_path)
        
        # Resize nếu ảnh quá lớn
        if img.width > max_width:
            ratio = max_width / img.width
            new_height = int(img.height * ratio)
            img = img.resize((max_width, new_height), Image.LANCZOS)
        
        # Lưu lại với chất lượng tối ưu
        img.save(image_path, optimize=True, quality=quality)
        
        return True
    except Exception as e:
        return False

def detect_table_in_image(image_path, language='vie+eng'):
    """Phát hiện và trích xuất bảng từ hình ảnh bằng OCR"""
    if not TESSERACT_AVAILABLE:
        return None
    
    try:
        # Đọc ảnh
        img = cv2.imread(image_path)
        if img is None:
            return None
        
        # Chuyển sang grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Áp dụng threshold để làm nổi bật đường viền
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
        
        # Phát hiện đường ngang và dọc (đặc trưng của bảng)
        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
        vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
        
        horizontal_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
        vertical_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
        
        # Kết hợp đường ngang và dọc
        table_mask = cv2.add(horizontal_lines, vertical_lines)
        
        # Nếu phát hiện đủ nhiều đường => có thể là bảng
        if cv2.countNonZero(table_mask) > 100:
            # Sử dụng OCR để đọc text với ngôn ngữ tiếng Việt + English
            ocr_data = pytesseract.image_to_string(img, lang=language)
            
            # Thử parse thành bảng
            lines = [line.strip() for line in ocr_data.split('\n') if line.strip()]
            if len(lines) >= 2:  # Ít nhất có header và 1 row
                # Tạo bảng Markdown
                return lines
        
        return None
    except Exception as e:
        return None

def lines_to_markdown_table(lines):
    """Chuyển đổi danh sách dòng thành bảng Markdown"""
    if not lines or len(lines) < 2:
        return None
    
    markdown_table = ""
    
    # Header
    header_parts = [p.strip() for p in lines[0].split() if p.strip()]
    if len(header_parts) > 0:
        markdown_table += "| " + " | ".join(header_parts) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(header_parts)) + " |\n"
        
        # Rows
        for line in lines[1:]:
            row_parts = [p.strip() for p in line.split() if p.strip()]
            if len(row_parts) > 0:
                # Đảm bảo số cột bằng header
                while len(row_parts) < len(header_parts):
                    row_parts.append("")
                markdown_table += "| " + " | ".join(row_parts[:len(header_parts)]) + " |\n"
    
    return markdown_table if markdown_table else None

def extract_images_from_pdf(pdf_path, output_folder, optimize_imgs=True, enable_ocr=True, ocr_lang='vie+eng'):
    """Trích xuất hình ảnh từ PDF"""
    doc = fitz.open(pdf_path)
    images = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            # Lưu hình ảnh
            image_name = f"image_page{page_num + 1}_{img_index + 1}.{image_ext}"
            image_path = os.path.join(output_folder, image_name)
            
            with open(image_path, "wb") as img_file:
                img_file.write(image_bytes)
            
            # Tối ưu hóa ảnh nếu được bật
            if optimize_imgs:
                optimize_image(image_path)
            
            # Thử phát hiện bảng trong ảnh nếu OCR được bật
            table_data = None
            if enable_ocr:
                table_data = detect_table_in_image(image_path, ocr_lang)
            
            images.append({
                'page': page_num,
                'path': image_path,
                'name': image_name,
                'is_table': table_data is not None,
                'table_data': table_data
            })
    
    doc.close()
    return images

def pdf_to_markdown(pdf_path, output_folder, image_path_prefix='', optimize_imgs=True, enable_ocr=True, ocr_lang='vie+eng'):
    """Chuyển đổi PDF sang Markdown"""
    doc = fitz.open(pdf_path)
    markdown_content = ""
    
    # Tạo thư mục cho hình ảnh
    os.makedirs(output_folder, exist_ok=True)
    
    # Trích xuất hình ảnh
    images = extract_images_from_pdf(pdf_path, output_folder, optimize_imgs, enable_ocr, ocr_lang)
    image_index = 0
    
    stats = {
        'pages': len(doc),
        'images': len(images),
        'tables': sum(1 for img in images if img.get('is_table', False))
    }
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Lấy văn bản
        text = page.get_text()
        
        # Thêm văn bản vào markdown
        if text.strip():
            # Phân tích cấu trúc cơ bản
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # Phát hiện tiêu đề (dòng ngắn, in hoa hoặc có font size lớn)
                    if len(line) < 100 and (line.isupper() or len(line.split()) <= 10):
                        markdown_content += f"\n## {line}\n\n"
                    else:
                        markdown_content += f"{line}\n\n"
        
        # Thêm hình ảnh từ trang này
        page_images = [img for img in images if img['page'] == page_num]
        for img in page_images:
            # Nếu ảnh là bảng, hiển thị bảng thay vì ảnh
            if img.get('is_table') and img.get('table_data'):
                table_md = lines_to_markdown_table(img['table_data'])
                if table_md:
                    markdown_content += f"\n**📊 Bảng (OCR):**\n\n{table_md}\n\n"
                else:
                    image_path = f"{image_path_prefix}{img['name']}"
                    markdown_content += f"![Image]({image_path})\n\n"
            else:
                image_path = f"{image_path_prefix}{img['name']}"
                markdown_content += f"![Image]({image_path})\n\n"
        
        # Phân cách trang
        if page_num < len(doc) - 1:
            markdown_content += "\n---\n\n"
    
    doc.close()
    return markdown_content, stats

def extract_images_from_docx(docx_path, output_folder):
    """Trích xuất hình ảnh từ Word"""
    doc = Document(docx_path)
    images = []
    
    os.makedirs(output_folder, exist_ok=True)
    
    # Trích xuất hình ảnh từ relationships
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_ext = rel.target_ref.split('.')[-1]
            image_name = f"image_{len(images) + 1}.{image_ext}"
            image_path = os.path.join(output_folder, image_name)
            
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)
            
            images.append({
                'path': image_path,
                'name': image_name
            })
    
    return images

def docx_to_markdown(docx_path, output_folder, image_path_prefix=''):
    """Chuyển đổi Word sang Markdown"""
    doc = Document(docx_path)
    markdown_content = ""
    
    # Tạo thư mục cho hình ảnh
    os.makedirs(output_folder, exist_ok=True)
    
    # Trích xuất hình ảnh
    images = extract_images_from_docx(docx_path, output_folder)
    image_index = 0
    
    stats = {
        'paragraphs': len(doc.paragraphs),
        'images': len(images),
        'tables': len(doc.tables)
    }
    
    for element in doc.element.body:
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, doc)
            text = paragraph.text.strip()
            
            if text:
                # Xác định style
                style = paragraph.style.name.lower()
                
                if 'heading 1' in style:
                    markdown_content += f"# {text}\n\n"
                elif 'heading 2' in style:
                    markdown_content += f"## {text}\n\n"
                elif 'heading 3' in style:
                    markdown_content += f"### {text}\n\n"
                elif 'heading 4' in style:
                    markdown_content += f"#### {text}\n\n"
                elif 'heading 5' in style:
                    markdown_content += f"##### {text}\n\n"
                elif 'heading 6' in style:
                    markdown_content += f"###### {text}\n\n"
                else:
                    # Xử lý định dạng văn bản
                    formatted_text = text
                    for run in paragraph.runs:
                        run_text = run.text
                        if run.bold:
                            formatted_text = formatted_text.replace(run_text, f"**{run_text}**")
                        if run.italic:
                            formatted_text = formatted_text.replace(run_text, f"*{run_text}*")
                    
                    markdown_content += f"{formatted_text}\n\n"
            
            # Kiểm tra xem paragraph có chứa hình ảnh không
            if paragraph._element.xpath('.//pic:pic'):
                if image_index < len(images):
                    image_path = f"{image_path_prefix}{images[image_index]['name']}"
                    markdown_content += f"![Image]({image_path})\n\n"
                    image_index += 1
        
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            markdown_content += "\n"
            
            # Header
            if table.rows:
                header_cells = table.rows[0].cells
                markdown_content += "| " + " | ".join([cell.text.strip() for cell in header_cells]) + " |\n"
                markdown_content += "| " + " | ".join(["---" for _ in header_cells]) + " |\n"
                
                # Các hàng còn lại
                for row in table.rows[1:]:
                    markdown_content += "| " + " | ".join([cell.text.strip() for cell in row.cells]) + " |\n"
            
            markdown_content += "\n"
    
    return markdown_content, stats

def create_zip_file(markdown_content, images_dir, md_filename):
    """Tạo file ZIP chứa markdown và tất cả hình ảnh"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Thêm file markdown
        zip_file.writestr(md_filename, markdown_content)
        
        # Thêm tất cả hình ảnh
        if os.path.exists(images_dir):
            for filename in os.listdir(images_dir):
                file_path = os.path.join(images_dir, filename)
                if os.path.isfile(file_path):
                    zip_file.write(file_path, f"images/{filename}")
    
    zip_buffer.seek(0)
    return zip_buffer

def main():
    st.set_page_config(page_title="Chuyển đổi PDF/Word sang Markdown", page_icon="📝", layout="wide")
    
    st.title("📝 Chuyển đổi PDF/Word sang Markdown")
    st.write("Upload file PDF hoặc Word để chuyển đổi sang định dạng Markdown (bao gồm cả hình ảnh)")
    
    # Hiển thị trạng thái OCR
    if TESSERACT_AVAILABLE:
        st.success("✅ OCR đã được kích hoạt - Có thể nhận diện bảng từ hình ảnh (Tiếng Việt + English)!")
    else:
        st.warning("⚠️ OCR chưa khả dụng - Bảng sẽ hiển thị dưới dạng hình ảnh")
    
    # Sidebar - Tùy chọn
    with st.sidebar:
        st.header("⚙️ Tùy chọn")
        
        st.subheader("📊 OCR Settings")
        enable_ocr = st.checkbox("Bật OCR nhận diện bảng", value=True, help="Tự động phát hiện và chuyển đổi bảng từ hình ảnh")
        ocr_language = st.selectbox(
            "Ngôn ngữ OCR",
            ["vie+eng", "eng", "vie"],
            help="vie+eng: Tiếng Việt + English (khuyên dùng)"
        )
        
        st.subheader("🖼️ Hình ảnh")
        optimize_images = st.checkbox("Tối ưu kích thước ảnh", value=True, help="Giảm kích thước ảnh để file nhẹ hơn")
        image_path = st.selectbox(
            "Đường dẫn ảnh trong Markdown",
            ["", "images/", "./", "./images/"],
            help="Chọn format đường dẫn ảnh trong file Markdown"
        )
        
        st.subheader("📦 Export")
        export_format = st.multiselect(
            "Format xuất file",
            ["Markdown (.md)", "ZIP (MD + Images)", "HTML"],
            default=["Markdown (.md)", "ZIP (MD + Images)"]
        )
    
    # Upload file (có thể nhiều file)
    uploaded_files = st.file_uploader(
        "Chọn file PDF hoặc Word", 
        type=['pdf', 'docx'],
        accept_multiple_files=True,
        help="Có thể chọn nhiều file cùng lúc"
    )
    
    if uploaded_files:
        # Hiển thị thông tin files
        if len(uploaded_files) == 1:
            st.info(f"📄 File: {uploaded_files[0].name} ({uploaded_files[0].size / 1024:.2f} KB)")
        else:
            st.info(f"📄 Đã chọn {len(uploaded_files)} files - Tổng: {sum(f.size for f in uploaded_files) / 1024:.2f} KB")
        
        # Nút chuyển đổi
        if st.button("🚀 Chuyển đổi sang Markdown", type="primary"):
            start_time = time.time()
            
            all_results = []
            
            for uploaded_file in uploaded_files:
                with st.spinner(f"Đang xử lý {uploaded_file.name}..."):
                    try:
                        # Tạo thư mục tạm để xử lý
                        temp_dir = f"temp_output_{uploaded_file.name.replace('.', '_')}"
                        images_dir = os.path.join(temp_dir, "images")
                        os.makedirs(images_dir, exist_ok=True)
                        
                        # Lưu file tạm
                        temp_file_path = os.path.join(temp_dir, uploaded_file.name)
                        with open(temp_file_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())
                        
                        # Xác định loại file và xử lý
                        if uploaded_file.name.endswith('.pdf'):
                            markdown_content, stats = pdf_to_markdown(
                                temp_file_path, images_dir, image_path, 
                                optimize_images, enable_ocr, ocr_language
                            )
                        elif uploaded_file.name.endswith('.docx'):
                            markdown_content, stats = docx_to_markdown(temp_file_path, images_dir, image_path)
                        else:
                            st.error(f"❌ {uploaded_file.name}: Định dạng không được hỗ trợ!")
                            continue
                        
                        all_results.append({
                            'filename': uploaded_file.name,
                            'markdown': markdown_content,
                            'stats': stats,
                            'images_dir': images_dir,
                            'temp_dir': temp_dir
                        })
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi khi xử lý {uploaded_file.name}: {str(e)}")
                        continue
            
            if all_results:
                elapsed_time = time.time() - start_time
                st.success(f"✅ Chuyển đổi thành công {len(all_results)} file(s) trong {elapsed_time:.2f}s!")
                
                # Nếu chỉ 1 file, hiển thị chi tiết
                if len(all_results) == 1:
                    result = all_results[0]
                    
                    # Hiển thị thống kê
                    st.subheader("📊 Thống kê")
                    stat_cols = st.columns(len(result['stats']))
                    for idx, (key, value) in enumerate(result['stats'].items()):
                        with stat_cols[idx]:
                            st.metric(key.capitalize(), value)
                    
                    # Tạo 2 cột
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.subheader("📝 Nội dung Markdown")
                        st.text_area("Markdown Output", result['markdown'], height=400)
                        
                        # Nút download markdown
                        if "Markdown (.md)" in export_format:
                            st.download_button(
                                label="💾 Tải xuống Markdown",
                                data=result['markdown'],
                                file_name=f"{Path(result['filename']).stem}.md",
                                mime="text/markdown"
                            )
                    
                        
                        # Nút download ZIP
                        if "ZIP (MD + Images)" in export_format:
                            zip_file = create_zip_file(
                                result['markdown'], 
                                result['images_dir'], 
                                f"{Path(result['filename']).stem}.md"
                            )
                            st.download_button(
                                label="📦 Tải xuống ZIP (MD + Images)",
                                data=zip_file,
                                file_name=f"{Path(result['filename']).stem}.zip",
                                mime="application/zip"
                            )
                        
                        # HTML Export
                        if "HTML" in export_format:
                            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{Path(result['filename']).stem}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; }}
        img {{ max-width: 100%; height: auto; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{result['markdown']}
</body>
</html>
"""
                            st.download_button(
                                label="📄 Tải xuống HTML",
                                data=html_content,
                                file_name=f"{Path(result['filename']).stem}.html",
                                mime="text/html"
                            )
                    
                    with col2:
                        st.subheader("👁️ Preview Markdown")
                        # Thay thế đường dẫn ảnh để hiển thị trong Streamlit
                        preview_content = result['markdown']
                        if os.path.exists(result['images_dir']):
                            image_files = [f for f in os.listdir(result['images_dir']) if f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
                            for img_file in image_files:
                                img_path = os.path.join(result['images_dir'], img_file)
                                # Đọc ảnh và convert sang base64 để hiển thị inline
                                with open(img_path, "rb") as img_f:
                                    img_data = base64.b64encode(img_f.read()).decode()
                                    img_ext = img_file.split('.')[-1]
                                    # Thay thế đường dẫn ảnh bằng data URI
                                    preview_content = preview_content.replace(
                                        f"![Image]({image_path}{img_file})",
                                        f'<img src="data:image/{img_ext};base64,{img_data}" alt="{img_file}" style="max-width:100%; height:auto;"/>'
                                    )
                                    preview_content = preview_content.replace(
                                        f"![Image]({img_file})",
                                        f'<img src="data:image/{img_ext};base64,{img_data}" alt="{img_file}" style="max-width:100%; height:auto;"/>'
                                    )
                        st.markdown(preview_content, unsafe_allow_html=True)
                    
                    # Hiển thị hình ảnh đã trích xuất
                    if os.path.exists(result['images_dir']):
                        image_files = [f for f in os.listdir(result['images_dir']) if f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
                        if image_files:
                            st.subheader(f"🖼️ Hình ảnh đã trích xuất ({len(image_files)} ảnh)")
                            
                            cols = st.columns(3)
                            for idx, img_file in enumerate(image_files):
                                with cols[idx % 3]:
                                    img_path = os.path.join(result['images_dir'], img_file)
                                    st.image(img_path, caption=img_file, use_container_width=True)
                                    
                                    # Nút download từng ảnh
                                    with open(img_path, "rb") as f:
                                        st.download_button(
                                            label=f"⬇️ Tải {img_file}",
                                            data=f,
                                            file_name=img_file,
                                            mime="image/png",
                                            key=f"download_{img_file}"
                                        )
                
                # Nếu nhiều file, hiển thị tổng hợp
                else:
                    st.subheader("📊 Tổng hợp kết quả")
                    
                    for idx, result in enumerate(all_results):
                        with st.expander(f"📄 {result['filename']}", expanded=False):
                            # Thống kê
                            stat_cols = st.columns(len(result['stats']))
                            for idx_stat, (key, value) in enumerate(result['stats'].items()):
                                with stat_cols[idx_stat]:
                                    st.metric(key.capitalize(), value)
                            
                            # Download buttons
                            cols = st.columns(3)
                            with cols[0]:
                                if "Markdown (.md)" in export_format:
                                    st.download_button(
                                        label="💾 MD",
                                        data=result['markdown'],
                                        file_name=f"{Path(result['filename']).stem}.md",
                                        mime="text/markdown",
                                        key=f"md_{idx}"
                                    )
                            with cols[1]:
                                if "ZIP (MD + Images)" in export_format:
                                    zip_file = create_zip_file(
                                        result['markdown'], 
                                        result['images_dir'], 
                                        f"{Path(result['filename']).stem}.md"
                                    )
                                    st.download_button(
                                        label="📦 ZIP",
                                        data=zip_file,
                                        file_name=f"{Path(result['filename']).stem}.zip",
                                        mime="application/zip",
                                        key=f"zip_{idx}"
                                    )
                            with cols[2]:
                                if "HTML" in export_format:
                                    html_content = f"<!DOCTYPE html><html><body>{result['markdown']}</body></html>"
                                    st.download_button(
                                        label="📄 HTML",
                                        data=html_content,
                                        file_name=f"{Path(result['filename']).stem}.html",
                                        mime="text/html",
                                        key=f"html_{idx}"
                                    )
                    
                    # Download tất cả thành 1 ZIP lớn
                    st.subheader("📦 Tải xuống tất cả")
                    all_zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(all_zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for result in all_results:
                            # Thêm markdown
                            zip_file.writestr(
                                f"{Path(result['filename']).stem}/{Path(result['filename']).stem}.md",
                                result['markdown']
                            )
                            # Thêm images
                            if os.path.exists(result['images_dir']):
                                for img_file in os.listdir(result['images_dir']):
                                    img_path = os.path.join(result['images_dir'], img_file)
                                    zip_file.write(
                                        img_path,
                                        f"{Path(result['filename']).stem}/images/{img_file}"
                                    )
                    all_zip_buffer.seek(0)
                    st.download_button(
                        label="📦 Tải xuống tất cả (ZIP)",
                        data=all_zip_buffer,
                        file_name=f"converted_files_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip"
                    )
        
        # Hướng dẫn
        with st.expander("ℹ️ Hướng dẫn sử dụng"):
            st.markdown("""
            ### Cách sử dụng:
            1. **Upload file**: Chọn file PDF hoặc Word (.docx) cần chuyển đổi
            2. **Nhấn nút Chuyển đổi**: Ứng dụng sẽ xử lý file và tạo Markdown
            3. **Xem kết quả**: 
               - Nội dung Markdown ở cột bên trái
               - Preview Markdown ở cột bên phải
               - Hình ảnh được hiển thị bên dưới
            4. **Tải xuống**: Tải file Markdown và hình ảnh về máy
            
            ### Tính năng MỚI:
            - ✅ **Upload nhiều file cùng lúc** - Batch processing
            - ✅ **OCR tiếng Việt + English** - Nhận diện bảng chính xác
            - ✅ **Tối ưu hình ảnh** - Giảm kích thước file
            - ✅ **Tải xuống ZIP** - Gộp markdown + images
            - ✅ **Export HTML** - Xuất sang định dạng HTML
            - ✅ **Thống kê file** - Số trang, ảnh, bảng
            - ✅ **Tùy chọn đường dẫn ảnh** - Linh hoạt theo nhu cầu
            - ✅ **PWA Ready** - Cài đặt như app native
            
            ### Tính năng cơ bản:
            - ✅ Chuyển đổi PDF sang Markdown
            - ✅ Chuyển đổi Word (.docx) sang Markdown
            - ✅ Trích xuất và lưu hình ảnh
            - ✅ Giữ nguyên định dạng cơ bản (tiêu đề, in đậm, in nghiêng)
            - ✅ Hỗ trợ bảng (từ Word)
            - ✅ Preview trực tiếp
            
            ### Lưu ý:
            - Hình ảnh sẽ được lưu trong thư mục `images/`
            - Đường dẫn hình ảnh trong Markdown là đường dẫn tương đối
            - Định dạng phức tạp có thể cần chỉnh sửa thủ công
            """)

if __name__ == "__main__":
    main()
